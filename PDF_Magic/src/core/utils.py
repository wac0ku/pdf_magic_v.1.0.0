# Autor: Leon Gajtner 
# Datum: 07.11.2024
# PDF Magic Utils
# Version: 2.1

import os
import logging
from datetime import datetime
from typing import List, Optional
import shutil

from PyPDF2 import PdfReader
from docx import Document
from PyQt5.QtWidgets import QTextEdit, QProgressBar, QMessageBox
from PyQt5.QtCore import QObject, pyqtSignal

class ConversionError(Exception):
    """Benutzerdefinierte Ausnahme für Konvertierungsfehler"""
    pass

class ConversionWorker(QObject):
    """Worker-Klasse für die PDF-zu-DOCX Konvertierung"""
    finished = pyqtSignal()
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, pdf_files: List[str], output_dir: str):
        super().__init__()
        self.pdf_files = pdf_files
        self.output_dir = output_dir
        self.setup_logging()

    def setup_logging(self):
        """Richtet das Logging für den Worker ein"""
        self.logger = logging.getLogger('ConversionWorker')
        self.logger.setLevel(logging.INFO)
        handler = logging.FileHandler('conversion.log')
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        self.logger.addHandler(handler)

    def run(self):
        """Führt die Konvertierung für alle PDF-Dateien durch"""
        total_files = len(self.pdf_files)
        successful_conversions = 0
        failed_conversions = 0

        for i, pdf_path in enumerate(self.pdf_files):
            try:
                if not os.path.exists(pdf_path):
                    raise FileNotFoundError(f"Die Datei {pdf_path} existiert nicht.")

                if not os.access(pdf_path, os.R_OK):
                    raise PermissionError(f"Keine Leserechte für {pdf_path}.")

                output_path = self.get_output_path(pdf_path)
                
                if os.path.exists(output_path):
                    new_path = self.get_unique_filename(output_path)
                    self.log.emit(f"Datei existiert bereits. Verwende neuen Namen: {os.path.basename(new_path)}")
                    output_path = new_path

                self.convert_pdf_to_docx(pdf_path, output_path)
                successful_conversions += 1

                progress = int(((i + 1) / total_files) * 100)
                self.progress.emit(progress)
                self.log.emit(f"Erfolgreich konvertiert: {pdf_path} -> {output_path}")

            except FileNotFoundError as e:
                failed_conversions += 1
                self.handle_error(str(e), pdf_path)
            except PermissionError as e:
                failed_conversions += 1
                self.handle_error(str(e), pdf_path)
            except ConversionError as e:
                failed_conversions += 1
                self.handle_error(str(e), pdf_path)
            except Exception as e:
                failed_conversions += 1
                self.handle_error(f"Unerwarteter Fehler: {str(e)}", pdf_path)

        self.log.emit(self.get_summary(successful_conversions, failed_conversions, total_files))
        self.finished.emit()

    def handle_error(self, error_msg: str, file_path: str):
        """Behandelt Fehler während der Konvertierung"""
        full_error_msg = f"Fehler bei der Konvertierung von {file_path}: {error_msg}"
        self.logger.error(full_error_msg)
        self.error.emit(full_error_msg)

    def convert_pdf_to_docx(self, pdf_path: str, docx_path: str):
        """Konvertiert eine PDF-Datei in eine DOCX-Datei."""
        try:
            pdf_reader = PdfReader(pdf_path)
            doc = Document()

            doc.core_properties.author = "PDF Magic"
            doc.core_properties.created = datetime.now()

            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text:
                    doc.add_heading(f'Seite {page_num}', level=1)
                    doc.add_paragraph(text)
                else:
                    self.log.emit(f"Warnung: Seite {page_num} in {pdf_path} enthält keinen extrahierbaren Text.")

            doc.save(docx_path)

        except Exception as e:
            raise ConversionError(f"Fehler bei der Konvertierung: {str(e)}")

    def get_output_path(self, pdf_path: str) -> str:
        """Erstellt den Ausgabepfad für die DOCX-Datei."""
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        return os.path.join(self.output_dir, f"{base_name}.docx")

    def get_unique_filename(self, file_path: str) -> str:
        """Erstellt einen eindeutigen Dateinamen wenn die Datei bereits existiert."""
        base, ext = os.path.splitext(file_path)
        counter = 1
        while os.path.exists(file_path):
            file_path = f"{base}_{counter}{ext}"
            counter += 1
        return file_path

    def get_summary(self, successful: int, failed: int, total: int) -> str:
        """Erstellt eine Zusammenfassung der Konvertierung."""
        return (f"Konvertierung abgeschlossen.\n"
                f"Erfolgreich: {successful}\n"
                f"Fehlgeschlagen: {failed}\n"
                f"Gesamt: {total}")

def update_log(log_window: QTextEdit, message: str):
    """Aktualisiert das Log-Fenster mit einer neuen Nachricht."""
    log_window.append(message)
    log_window.verticalScrollBar().setValue(log_window.verticalScrollBar().maximum())

def update_progress_bar(progress_bar: QProgressBar, value: int):
    """Aktualisiert den Fortschrittsbalken."""
    progress_bar.setValue(value)

def show_error_message(message: str):
    """Zeigt eine Fehlermeldung in einem Dialogfenster an."""
    error_box = QMessageBox()
    error_box.setIcon(QMessageBox.Critical)
    error_box.setText("Ein Fehler ist aufgetreten")
    error_box.setInformativeText(message)
    error_box.setWindowTitle("Fehler")
    error_box.exec_()

def validate_pdf_file(file_path: str) -> Optional[str]:
    """
    Überprüft, ob die angegebene Datei eine gültige PDF-Datei ist.
    
    Args:
        file_path: Pfad zur zu überprüfenden Datei
        
    Returns:
        Optional[str]: Fehlermeldung wenn die Datei ungültig ist, sonst None
    """
    try:
        if not os.path.exists(file_path):
            return f"Die Datei {file_path} existiert nicht."
        
        if not os.access(file_path, os.R_OK):
            return f"Keine Leserechte für {file_path}."
        
        if not file_path.lower().endswith('.pdf'):
            return f"Die Datei {file_path} ist keine PDF-Datei."
        
        # Überprüfe die Dateigröße (max. 100MB)
        file_size = os.path.getsize(file_path)
        if file_size > 100 * 1024 * 1024:  # 100MB in Bytes
            return f"Die Datei {file_path} ist zu groß (max. 100MB erlaubt)."
        
        # Versuche die PDF-Datei zu öffnen um ihre Gültigkeit zu prüfen
        with open(file_path, 'rb') as file:
            try:
                PdfReader(file)
            except Exception:
                return f"Die Datei {file_path} ist keine gültige PDF-Datei."
        
        return None
        
    except Exception as e:
        return f"Fehler bei der Überprüfung von {file_path}: {str(e)}"

def validate_output_directory(directory: str) -> Optional[str]:
    """
    Überprüft, ob das Ausgabeverzeichnis gültig und beschreibbar ist.
    
    Args:
        directory: Pfad zum zu überprüfenden Verzeichnis
        
    Returns:
        Optional[str]: Fehlermeldung wenn das Verzeichnis ungültig ist, sonst None
    """
    try:
        if not directory:
            return "Kein Ausgabeverzeichnis angegeben."
        
        if not os.path.exists(directory):
            try:
                os.makedirs(directory)
            except Exception as e:
                return f"Konnte Ausgabeverzeichnis nicht erstellen: {str(e)}"
        
        if not os.access(directory, os.W_OK):
            return f"Keine Schreibrechte im Verzeichnis {directory}."
        
        return None
        
    except Exception as e:
        return f"Fehler bei der Überprüfung des Verzeichnisses: {str(e)}"

def check_disk_space(directory: str, required_space: int) -> Optional[str]:
    """
    Überprüft, ob genügend Speicherplatz verfügbar ist.
    
    Args:
        directory: Zu überprüfendes Verzeichnis
        required_space: Benötigter Speicherplatz in Bytes
        
    Returns:
        Optional[str]: Fehlermeldung wenn nicht genug Speicherplatz vorhanden ist, sonst None
    """
    try:
        total, used, free = shutil.disk_usage(directory)
        if free < required_space:
            free_mb = free / (1024 * 1024)
            required_mb = required_space / (1024 * 1024)
            return (f"Nicht genügend Speicherplatz verfügbar. "
                   f"Verfügbar: {free_mb:.1f}MB, Benötigt: {required_mb:.1f}MB")
        return None
    except Exception as e:
        return f"Fehler bei der Überprüfung des Speicherplatzes: {str(e)}"

def sanitize_filename(filename: str) -> str:
    """
    Bereinigt einen Dateinamen von ungültigen Zeichen.
    
    Args:
        filename: Zu bereinigender Dateiname
        
    Returns:
        str: Bereinigter Dateiname
    """
    # Ungültige Zeichen für Dateinamen
    invalid_chars = '<>:"/\\|?*'
    # Ersetze ungültige Zeichen durch Unterstrich
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    # Entferne führende und nachfolgende Leerzeichen
    filename = filename.strip()
    # Begrenze die Länge des Dateinamens
    if len(filename) > 255:
        base, ext = os.path.splitext(filename)
        filename = base[:255-len(ext)] + ext
    return filename

def create_backup(file_path: str) -> Optional[str]:
    """
    Erstellt eine Sicherungskopie einer Datei.
    
    Args:
        file_path: Pfad zur zu sichernden Datei
        
    Returns:
        Optional[str]: Pfad zur Backup-Datei oder None bei Fehler
    """
    try:
        if not os.path.exists(file_path):
            return None
            
        backup_dir = os.path.join(os.path.dirname(file_path), 'backups')
        if not os.path.exists(backup_dir):
            os.makedirs(backup_dir)
            
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_name = os.path.basename(file_path)
        backup_path = os.path.join(backup_dir, f"{timestamp}_{base_name}")
        
        shutil.copy2(file_path, backup_path)
        return backup_path
        
    except Exception as e:
        logging.error(f"Fehler beim Erstellen des Backups: {str(e)}")
        return None

class FileProcessor:
    """Klasse zur Verarbeitung von Dateien mit Fehlerbehandlung"""
    
    def __init__(self):
        self.logger = logging.getLogger('FileProcessor')
        
    def process_file(self, file_path: str, output_dir: str) -> bool:
        """
        Verarbeitet eine einzelne Datei mit umfassender Fehlerbehandlung.
        
        Args:
            file_path: Pfad zur Eingabedatei
            output_dir: Ausgabeverzeichnis
            
        Returns:
            bool: True wenn erfolgreich, False bei Fehler
        """
        try:
            # Validiere Eingabedatei
            error_msg = validate_pdf_file(file_path)
            if error_msg:
                self.logger.error(error_msg)
                return False
                
            # Validiere Ausgabeverzeichnis
            error_msg = validate_output_directory(output_dir)
            if error_msg:
                self.logger.error(error_msg)
                return False
                
            # Überprüfe Speicherplatz
            file_size = os.path.getsize(file_path)
            error_msg = check_disk_space(output_dir, file_size * 2)  # Schätze doppelte Größe
            if error_msg:
                self.logger.error(error_msg)
                return False
                
            # Erstelle Backup
            backup_path = create_backup(file_path)
            if backup_path:
                self.logger.info(f"Backup erstellt: {backup_path}")
                
            return True
            
        except Exception as e:
            self.logger.error(f"Unerwarteter Fehler bei der Verarbeitung von {file_path}: {str(e)}")
            return False

def cleanup_temp_files(directory: str, prefix: str = "temp_", older_than: int = 24):
    """
    Bereinigt temporäre Dateien in einem bestimmten Verzeichnis.

    Args:
        directory (str): Das zu bereinigende Verzeichnis
        prefix (str): Präfix der zu löschenden temporären Dateien
        older_than (int): Lösche Dateien, die älter als diese Anzahl von Stunden sind

    Returns:
        int: Anzahl der gelöschten Dateien
    """
    deleted_count = 0
    current_time = datetime.now()
    try:
        for filename in os.listdir(directory):
            if filename.startswith(prefix):
                file_path = os.path.join(directory, filename)
                file_modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
                if (current_time - file_modified_time).total_seconds() > older_than * 3600:
                    os.remove(file_path)
                    deleted_count += 1
                    logging.info(f"Gelöschte temporäre Datei: {file_path}")
    except Exception as e:
        logging.error(f"Fehler beim Bereinigen temporärer Dateien: {str(e)}")
    return deleted_count

def get_file_info(file_path: str) -> dict:
    """
    Gibt Informationen über eine Datei zurück.

    Args:
        file_path (str): Pfad zur Datei

    Returns:
        dict: Ein Dictionary mit Dateiinformationen
    """
    try:
        stat_info = os.stat(file_path)
        return {
            "name": os.path.basename(file_path),
            "size": stat_info.st_size,
            "created": datetime.fromtimestamp(stat_info.st_ctime),
            "modified": datetime.fromtimestamp(stat_info.st_mtime),
            "accessed": datetime.fromtimestamp(stat_info.st_atime),
        }
    except Exception as e:
        logging.error(f"Fehler beim Abrufen von Dateiinformationen für {file_path}: {str(e)}")
        return {}

def is_pdf_encrypted(file_path: str) -> bool:
    """
    Überprüft, ob eine PDF-Datei verschlüsselt ist.

    Args:
        file_path (str): Pfad zur PDF-Datei

    Returns:
        bool: True, wenn die PDF verschlüsselt ist, sonst False
    """
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            return pdf_reader.is_encrypted
    except Exception as e:
        logging.error(f"Fehler beim Überprüfen der PDF-Verschlüsselung für {file_path}: {str(e)}")
        return False

def estimate_conversion_time(file_size: int) -> float:
    """
    Schätzt die Konvertierungszeit basierend auf der Dateigröße.

    Args:
        file_size (int): Größe der Datei in Bytes

    Returns:
        float: Geschätzte Konvertierungszeit in Sekunden
    """
    # Diese Formel sollte an Ihre spezifischen Leistungsmerkmale angepasst werden
    base_time = 2  # Basiszeit in Sekunden
    time_per_mb = 0.5  # Zusätzliche Zeit pro MB
    estimated_time = base_time + (file_size / (1024 * 1024)) * time_per_mb
    return estimated_time

def create_error_report(error_log: List[str], output_path: str):
    """
    Erstellt einen Fehlerbericht basierend auf den gesammelten Fehlerprotokollen.

    Args:
        error_log (List[str]): Liste der Fehlerprotokolle
        output_path (str): Pfad, an dem der Fehlerbericht gespeichert werden soll
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as report_file:
            report_file.write("PDF Magic - Fehlerbericht\n")
            report_file.write("=" * 30 + "\n\n")
            report_file.write(f"Erstellungszeitpunkt: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            for error in error_log:
                report_file.write(f"{error}\n")
        logging.info(f"Fehlerbericht erstellt: {output_path}")
    except Exception as e:
        logging.error(f"Fehler beim Erstellen des Fehlerberichts: {str(e)}")

def verify_conversion(pdf_path: str, docx_path: str) -> bool:
    """
    Überprüft, ob die Konvertierung erfolgreich war, indem grundlegende Eigenschaften verglichen werden.

    Args:
        pdf_path (str): Pfad zur ursprünglichen PDF-Datei
        docx_path (str): Pfad zur konvertierten DOCX-Datei

    Returns:
        bool: True, wenn die Konvertierung als erfolgreich verifiziert wurde, sonst False
    """
    try:
        # Überprüfe, ob die DOCX-Datei existiert und nicht leer ist
        if not os.path.exists(docx_path) or os.path.getsize(docx_path) == 0:
            logging.error(f"Konvertierung fehlgeschlagen: {docx_path} existiert nicht oder ist leer.")
            return False

        # Vergleiche die Anzahl der Seiten/Absätze (grobe Schätzung)
        pdf_reader = PdfReader(pdf_path)
        pdf_pages = len(pdf_reader.pages)

        doc = Document(docx_path)
        docx_paragraphs = len(doc.paragraphs)

        # Einfache Heuristik: Anzahl der DOCX-Absätze sollte mindestens der Seitenanzahl der PDF entsprechen
        if docx_paragraphs < pdf_pages:
            logging.warning(f"Mögliche unvollständige Konvertierung: PDF hat {pdf_pages} Seiten, "
                            f"aber DOCX hat nur {docx_paragraphs} Absätze.")
            return False

        logging.info(f"Konvertierung erfolgreich verifiziert: {pdf_path} -> {docx_path}")
        return True

    except Exception as e:
        logging.error(f"Fehler bei der Verifizierung der Konvertierung: {str(e)}")
        return False