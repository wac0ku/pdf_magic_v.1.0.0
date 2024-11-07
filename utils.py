# Autor: Leon Gajtner 
# Datum: 07.11.2024
# PDF Magic Utils
# Version: 1.0.0

import os
import logging
from datetime import datetime
from PyPDF2 import PdfReader
from docx import Document
from PyQt5.QtWidgets import QTextEdit, QProgressBar, QMessageBox
from PyQt5.QtCore import QObject, pyqtSignal

class ConversionWorker(QObject):
    """Worker-Klasse für die PDF-zu-DOCX Konvertierung"""
    finished = pyqtSignal()
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, pdf_files, output_dir):
        """
        Initialisiert den ConversionWorker.
        
        Args:
            pdf_files (list): Liste der zu konvertierenden PDF-Dateipfade
            output_dir (str): Ausgabeverzeichnis für die DOCX-Dateien
        """
        super().__init__()
        self.pdf_files = pdf_files
        self.output_dir = output_dir
        self.setup_logging()

    def setup_logging(self):
        """Richtet das Logging für den Worker ein"""
        self.logger = logging.getLogger('ConversionWorker')
        self.logger.setLevel(logging.INFO)

    def run(self):
        """Führt die Konvertierung für alle PDF-Dateien durch"""
        total_files = len(self.pdf_files)
        successful_conversions = 0
        failed_conversions = 0

        for i, pdf_path in enumerate(self.pdf_files):
            try:
                # Erstelle den Ausgabepfad
                output_path = self.get_output_path(pdf_path)
                
                # Überprüfe ob die Ausgabedatei bereits existiert
                if os.path.exists(output_path):
                    new_path = self.get_unique_filename(output_path)
                    self.log.emit(f"Datei existiert bereits. Verwende neuen Namen: {os.path.basename(new_path)}")
                    output_path = new_path

                # Führe die Konvertierung durch
                self.convert_pdf_to_docx(pdf_path, output_path)
                successful_conversions += 1

                # Aktualisiere den Fortschritt
                progress = int(((i + 1) / total_files) * 100)
                self.progress.emit(progress)
                self.log.emit(f"Erfolgreich konvertiert: {pdf_path} -> {output_path}")

            except Exception as e:
                failed_conversions += 1
                error_msg = f"Fehler bei der Konvertierung von {pdf_path}: {str(e)}"
                self.logger.error(error_msg)
                self.error.emit(error_msg)

        # Sende abschließende Statistik
        summary = (f"Konvertierung abgeschlossen.\n"
                  f"Erfolgreich: {successful_conversions}\n"
                  f"Fehlgeschlagen: {failed_conversions}\n"
                  f"Gesamt: {total_files}")
        self.log.emit(summary)
        self.finished.emit()

    def convert_pdf_to_docx(self, pdf_path, docx_path):
        """
        Konvertiert eine PDF-Datei in eine DOCX-Datei.
        
        Args:
            pdf_path (str): Pfad zur PDF-Datei
            docx_path (str): Pfad für die zu erstellende DOCX-Datei
        """
        try:
            pdf_reader = PdfReader(pdf_path)
            doc = Document()

            # Füge Metadaten hinzu
            doc.core_properties.author = "PDF Magic"
            doc.core_properties.created = datetime.now()

            # Extrahiere Text von jeder Seite
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text:
                    # Füge Seitenüberschrift hinzu
                    doc.add_heading(f'Seite {page_num}', level=1)
                    doc.add_paragraph(text)
                else:
                    self.log.emit(f"Warnung: Seite {page_num} in {pdf_path} enthält keinen extrahierbaren Text.")

            # Speichere das Dokument
            doc.save(docx_path)

        except Exception as e:
            raise Exception(f"Fehler bei der Konvertierung: {str(e)}")

    def get_output_path(self, pdf_path):
        """
        Erstellt den Ausgabepfad für die DOCX-Datei.
        
        Args:
            pdf_path (str): Pfad zur PDF-Datei
            
        Returns:
            str: Pfad für die DOCX-Datei
        """
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        return os.path.join(self.output_dir, f"{base_name}.docx")

    def get_unique_filename(self, file_path):
        """
        Erstellt einen eindeutigen Dateinamen wenn die Datei bereits existiert.
        
        Args:
            file_path (str): Ursprünglicher Dateipfad
            
        Returns:
            str: Eindeutiger Dateipfad
        """
        base, ext = os.path.splitext(file_path)
        counter = 1
        while os.path.exists(file_path):
            file_path = f"{base}_{counter}{ext}"
            counter += 1
        return file_path
    
def update_log(log_window, message):
    """Aktualisiert das Log-Fenster mit einer neuen Nachricht."""
    log_window.append(message)
    log_window.verticalScrollBar().setValue(log_window.verticalScrollBar().maximum())

def update_progress_bar(progress_bar, value):
    """Aktualisiert den Fortschrittsbalken."""
    progress_bar.setValue(value)

def show_error_message(message):
    """Zeigt eine Fehlermeldung in einem Dialogfenster an."""
    error_box = QMessageBox()
    error_box.setIcon(QMessageBox.Critical)
    error_box.setText("Ein Fehler ist aufgetreten")
    error_box.setInformativeText(message)
    error_box.setWindowTitle("Fehler")
    error_box.exec_()