# converter.py
# Erstellt durch populate_files.py

import os
import logging
from datetime import datetime
from typing import List
from PyPDF2 import PdfReader
from docx import Document
from PyQt5.QtCore import QObject, pyqtSignal

class ConversionWorker(QObject):
    """Worker-Klasse für die PDF-zu-DOCX Konvertierung"""
    finished = pyqtSignal()
    progress = pyqtSignal(int)
    log = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, pdf_files: List[str], output_dir: str):
        super().__init__()
        self.pdf_files = pdf_files
        self.output_dir = output_dir
        self.setup_logging()

    def setup_logging(self):
        """Richtet das Logging für den Worker ein"""
        self.logger = logging.getLogger('ConversionWorker')
        self.logger.setLevel(logging.INFO)
        handler = logging.FileHandler('conversion.log')
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)
        self.logger.addHandler(handler)

    def run(self):
        """Führt die Konvertierung für alle PDF-Dateien durch"""
        total_files = len(self.pdf_files)
        successful_conversions = 0
        failed_conversions = 0

        for i, pdf_path in enumerate(self.pdf_files):
            try:
                if not os.path.exists(pdf_path):
                    raise FileNotFoundError(f"Die Datei {pdf_path} existiert nicht.")

                if not os.access(pdf_path, os.R_OK):
                    raise PermissionError(f"Keine Leserechte für {pdf_path}.")

                output_path = self.get_output_path(pdf_path)
                
                if os.path.exists(output_path):
                    new_path = self.get_unique_filename(output_path)
                    self.log.emit(f"Datei existiert bereits. Verwende neuen Namen: {os.path.basename(new_path)}")
                    output_path = new_path

                self.convert_pdf_to_docx(pdf_path, output_path)
                successful_conversions += 1

                progress = int(((i + 1) / total_files) * 100)
                self.progress.emit(progress)
                self.log.emit(f"Erfolgreich konvertiert: {pdf_path} -> {output_path}")

            except FileNotFoundError as e:
                failed_conversions += 1
                self.handle_error(str(e), pdf_path)
            except PermissionError as e:
                failed_conversions += 1
                self.handle_error(str(e), pdf_path)
            except ConversionError as e:
                failed_conversions += 1
                self.handle_error(str(e), pdf_path)
            except Exception as e:
                failed_conversions += 1
                self.handle_error(f"Unerwarteter Fehler: {str(e)}", pdf_path)

        self.log.emit(self.get_summary(successful_conversions, failed_conversions, total_files))
        self.finished.emit()

    def handle_error(self, error_msg: str, file_path: str):
        """Behandelt Fehler während der Konvertierung"""
        full_error_msg = f"Fehler bei der Konvertierung von {file_path}: {error_msg}"
        self.logger.error(full_error_msg)
        self.error.emit(full_error_msg)

    def convert_pdf_to_docx(self, pdf_path: str, docx_path: str):
        """Konvertiert eine PDF-Datei in eine DOCX-Datei."""
        try:
            pdf_reader = PdfReader(pdf_path)
            doc = Document()

            doc.core_properties.author = "PDF Magic"
            doc.core_properties.created = datetime.now()

            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text:
                    doc.add_heading(f'Seite {page_num}', level=1)
                    doc.add_paragraph(text)
                else:
                    self.log.emit(f"Warnung: Seite {page_num} in {pdf_path} enthält keinen extrahierbaren Text.")

            doc.save(docx_path)

        except Exception as e:
            raise ConversionError(f"Fehler bei der Konvertierung: {str(e)}")

    def get_output_path(self, pdf_path: str) -> str:
        """Erstellt den Ausgabepfad für die DOCX-Datei."""
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        return os.path.join(self.output_dir, f"{base_name}.docx")

    def get_unique_filename(self, file_path: str) -> str:
        """Erstellt einen eindeutigen Dateinamen wenn die Datei bereits existiert."""
        base, ext = os.path.splitext(file_path)
        counter = 1
        while os.path.exists(file_path):
            file_path = f"{base}_{counter}{ext}"
            counter += 1
        return file_path

    def get_summary(self, successful: int, failed: int, total: int) -> str:
        """Erstellt eine Zusammenfassung der Konvertierung."""
        return (f"Konvertierung abgeschlossen.\n"
                f"Erfolgreich: {successful}\n"
                f"Fehlgeschlagen: {failed}\n"
                f"Gesamt: {total}")